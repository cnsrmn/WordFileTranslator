# testing remote push
import os
import tkinter as tk
import shutil
import logging
import time
import threading
from tkinter import filedialog, messagebox, ttk
from googletrans import LANGUAGES, Translator
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
handler = logging.FileHandler("WordFileTranslator.log", delay=True)
formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
handler.setFormatter(formatter)
logger.addHandler(handler)


class ProgressWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.geometry('300x100')
        self.title("Working...")
        self.progress = ttk.Progressbar(self, length=200, mode='determinate')
        self.progress.pack(padx=10, pady=10)
        self.progress_label = tk.Label(self, text="0%")
        self.progress_label.pack()
        self.cancel_button = tk.Button(self, text="Cancel", command=self.cancel)
        self.cancel_button.pack(padx=10, pady=10)
        self.protocol("WM_DELETE_WINDOW", self.cancel)
        self.protocol("WM_DELETE_WINDOW", self.disable_close)
        self.cancelled = False
        self.grab_set()

    def disable_close(self):
        pass

    def run(self, func):
        self.thread = threading.Thread(target=func)
        self.thread.start()
        self.check_thread()

    def check_thread(self):
        if self.thread.is_alive():
            self.after(1000, self.check_thread)
        else:
            self.destroy()

    def cancel(self):
        self.cancelled = True


def translate_docx(doc_path, src_lang, dest_lang, progress_win):
    if not os.path.isfile(doc_path):
        messagebox.showwarning("Invalid File", f"The file {doc_path} does not exist.")
        return None
    if os.path.getsize(doc_path) == 0:
        messagebox.showwarning("Empty File", f"The file {doc_path} is empty and cannot be translated.")
        return None
    translator = Translator()
    try:
        doc = Document(doc_path)
    except (IOError, PackageNotFoundError) as e:
        logger.exception(e)
        messagebox.showerror("File Error", f"An error occurred while opening {doc_path}: {e}")
        return None
    file_name, file_ext = os.path.splitext(doc_path)
    dest_folder = os.path.dirname(doc_path)
    if not os.path.exists(dest_folder):
        os.makedirs(dest_folder)
    translated_doc_path = f"{file_name}_{dest_lang}{file_ext}"
    shutil.copy(doc_path, translated_doc_path)
    translated_doc = Document(translated_doc_path)
    start_time = time.time()
    chunk_size = 1024
    total_paragraphs = len(translated_doc.paragraphs)
    translated_paragraphs = 0
    for paragraph in translated_doc.paragraphs:
        if progress_win.cancelled:
            break
        text = paragraph.text
        translated_text = ""
        while text:
            if progress_win.cancelled:
                break
            chunk, text = text[:chunk_size], text[chunk_size:]
            try:
                translated_chunk = translator.translate(chunk, src=src_lang, dest=dest_lang).text
                translated_text += translated_chunk
            except Exception as e:
                logger.exception(e)
                num_chars = len(paragraph.text)
                num_words = len(paragraph.text.split())
                num_sentences = len(paragraph.text.split("."))
                logger.info(
                    f"Translated {num_chars} characters, {num_words} words, and {num_sentences} sentences from {src_lang} to {dest_lang}.")
            paragraph.text = translated_text
        translated_paragraphs += 1
        progress = (translated_paragraphs / total_paragraphs) * 100
        progress_win.progress['value'] = progress
        progress_win.progress_label['text'] = f"{progress:.2f}%"
        progress_win.update_idletasks()
    if not progress_win.cancelled:
        try:
            translated_doc.save(translated_doc_path)
        except IOError as e:
            logger.exception(e)
            messagebox.showerror("File Error", f"An error occurred while saving {translated_doc_path}: {e}")
            logger.exception(f"An error occurred while saving {translated_doc_path}: {e}")
            return None
    logging.shutdown()
    return translated_doc_path


def main():
    root = tk.Tk()
    root.geometry('800x400')
    root.title("WordFileTranslator")
    menubar = tk.Menu(root)
    root.config(menu=menubar)
    help_menu = tk.Menu(menubar, tearoff=False)
    menubar.add_cascade(label="Help", menu=help_menu)

    def show_about():
        tk.messagebox.showinfo("About",
                               "This is a program by Can SARMAN that allows you to translate Word documents from one language to another with ease and accuracy.\n\nIt uses the Google Translate API.\n\nWord File Translator version 1.0\n\nCurrent as of October 15, 2023")

    help_menu.add_command(label="About", command=show_about)
    frame = tk.Frame(root, padx=20, pady=10)
    frame.pack()

    def select_file():
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not file_path:
            return
        file_ext = os.path.splitext(file_path)[1]
        if file_ext != ".docx":
            messagebox.showwarning("Invalid File",
                                   f"The file {file_path} is not a Word document and cannot be translated.")
            return
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            messagebox.showwarning("Empty File", f"The file {file_path} is empty and cannot be translated.")
            return
        try:
            Document(file_path)
        except (IOError, PackageNotFoundError) as e:
            logger.exception(e)
            messagebox.showerror("File Error", f"An error occurred while opening {file_path}")
            return
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        translate_button.config(state='normal')

    select_file_button = tk.Button(frame, text="Select Word File To Be Translated", command=select_file)
    select_file_button.grid(row=0, column=0, columnspan=4, sticky='w', pady=10)
    source_file_label = tk.Label(frame, text="Source File", font=("Arial", 12))
    source_file_label.grid(row=1, column=0, sticky='w', pady=10)
    file_entry = tk.Entry(frame, width=100)
    file_entry.grid(row=1, column=1, columnspan=3, sticky='w', pady=10)
    lang_frame = tk.Frame(frame)
    lang_frame.grid(row=2, column=0, columnspan=4, sticky='w', pady=10)
    src_lang_label = tk.Label(lang_frame, text="Source Language")
    src_lang_label.grid(row=0, column=0)
    lang_frame.grid_columnconfigure(1, minsize=5)
    src_lang_box = ttk.Combobox(lang_frame, values=[lang.title() for lang in LANGUAGES.values()], width=20)
    src_lang_box.grid(row=0, column=2)
    src_lang_box.set('Turkish')
    lang_frame.grid_columnconfigure(3, minsize=5)
    dest_lang_label = tk.Label(lang_frame, text="Destination Language")
    dest_lang_label.grid(row=0, column=4)
    lang_frame.grid_columnconfigure(5, minsize=5)
    dest_lang_box = ttk.Combobox(lang_frame, values=[lang.title() for lang in LANGUAGES.values()], width=20)
    dest_lang_box.grid(row=0, column=6)
    dest_lang_box.set('English')
    lang_frame.grid_columnconfigure(7, minsize=10)
    # Disable the user from writing text into the combo boxes.
    src_lang_box.config(state='readonly')
    dest_lang_box.config(state='readonly')

    def translate():
        src_lang = src_lang_box.get()
        dest_lang = dest_lang_box.get()
        file_path = file_entry.get()
        if src_lang == dest_lang:
            messagebox.showwarning("Same Language",
                                   f"The source and destination languages are the same. Please select different languages.")
            return
        progress_win = ProgressWindow(root)

        def translate_with_progress():
            translated_doc_path = translate_docx(file_path, src_lang, dest_lang, progress_win)
            if translated_doc_path == file_path or progress_win.cancelled:
                file_entry.delete(0, tk.END)
                result_entry.delete(0, tk.END)
                return
            result_entry.delete(0, tk.END)
            result_entry.insert(0, os.path.abspath(translated_doc_path))
            log_file_entry.delete(0, tk.END)
            log_file_entry.insert(0, os.path.abspath("WordFileTranslator.log"))

            if not progress_win.cancelled and messagebox.askyesno("Translation Complete",
                                                                  f"The translated document can be found at {os.path.abspath(translated_doc_path)}. Do you want to open the folder?"):
                os.startfile(os.path.dirname(translated_doc_path))

        progress_win.run(translate_with_progress)

    translate_button = tk.Button(lang_frame, text="Translate", command=translate, font=("Arial", 18, "bold"), width=15)
    translate_button.config(state='disabled')
    translate_button.grid(row=0, column=8)
    translated_file_label = tk.Label(frame, text="Translated File", font=("Arial", 12))
    translated_file_label.grid(row=3, column=0, sticky='w', pady=10)
    log_file_label = tk.Label(frame, text="Log File", font=("Arial", 12))
    log_file_label.grid(row=4, column=0, sticky='w', pady=10)
    result_entry = tk.Entry(frame, width=100)
    result_entry.grid(row=3, column=1, columnspan=3, sticky='w', pady=10)
    log_file_entry = tk.Entry(frame, width=100)
    log_file_entry.grid(row=4, column=1, columnspan=3, sticky='w', pady=10)
    root.mainloop()


if __name__ == "__main__":
    main()